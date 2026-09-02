#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera SOP_PROCESSAMENTO_OTC.docx a partir de SOP_PROCESSAMENTO_OTC.md.

O Markdown é a FONTE ÚNICA do SOP. Edite o .md (inclua módulos, troque telas,
ajuste textos) e rode este script para regenerar o Word:

    pip install python-docx        # uma vez
    python scripts/build_sop_docx.py

Suporta: títulos (#..####), parágrafos, listas (- e 1.), tabelas |...|,
citações (>), **negrito**, `código`, imagens ![alt](caminho) e --- (régua).
Blocos <!-- ... --> do Markdown são ignorados (servem de comentário/modelo).
Caminhos de imagem são relativos à raiz do repositório.
"""
import os
import re
import sys
import atexit
import tempfile

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit('Falta a dependência: rode  pip install python-docx')

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Sem argumentos, gera o SOP (comportamento de sempre). Com argumentos, converte
# qualquer Markdown do repositório — é assim que o Guia do Usuário é gerado, sem
# precisar de um segundo conversor:
#     python scripts/build_sop_docx.py GUIA_DO_USUARIO_OTC_TRACKER.md
_ARGS = [a for a in sys.argv[1:] if not a.startswith('-')]
SRC = os.path.join(ROOT, _ARGS[0]) if _ARGS else os.path.join(ROOT, 'SOP_PROCESSAMENTO_OTC.md')
OUT = (os.path.join(ROOT, _ARGS[1]) if len(_ARGS) > 1
       else os.path.splitext(SRC)[0] + '.docx')
BLUE = RGBColor(0x00, 0x66, 0xCC)
GREY = RGBColor(0x55, 0x55, 0x55)
IMG_W = Inches(6.6)
_TMP_IMGS = []


@atexit.register
def _limpa_tmp():
    for f in _TMP_IMGS:
        try:
            os.unlink(f)
        except OSError:
            pass


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexc)
    tcPr.append(shd)


# As capturas são feitas em 2x (device_scale_factor=2 do Playwright), porque no
# Markdown elas são lidas na tela e o zoom precisa aguentar. No Word a imagem é
# desenhada com 6,6 polegadas de largura, então acima de ~1400 px o arquivo
# cresce sem ninguém enxergar diferença — o Guia saía com 44 MB, grande demais
# para anexar num e-mail. Aqui a cópia embutida é reduzida a essa largura; o PNG
# do repositório NÃO é tocado. Sem Pillow, embute o original (o documento sai
# gordo, mas sai).
IMG_MAX_PX = int(os.environ.get('SOP_IMG_MAX_PX') or 1400)
_shrunk = {}


def shrink(path):
    """Devolve o caminho de uma cópia reduzida da imagem, ou o próprio caminho."""
    if path in _shrunk:
        return _shrunk[path]
    out = path
    try:
        from PIL import Image
        im = Image.open(path)
        if im.width > IMG_MAX_PX:
            alt = int(round(im.height * IMG_MAX_PX / float(im.width)))
            im = im.convert('RGB').resize((IMG_MAX_PX, alt), Image.LANCZOS)
            # A cópia embutida é QUANTIZADA (256 cores + dither), como os PNGs
            # do repositório: o convert('RGB') acima joga a paleta fora, e as
            # capturas DARK em RGB comprimem tão mal que o Guia voltava aos
            # 55 MB mesmo com os arquivos do repo otimizados.
            im = im.quantize(256, method=Image.FASTOCTREE,
                             dither=Image.FLOYDSTEINBERG)
            tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            im.save(tmp.name, 'PNG', optimize=True)
            out = tmp.name
            _TMP_IMGS.append(tmp.name)
    except Exception:
        pass
    _shrunk[path] = out
    return out


def add_inline(par, text):
    for tok in re.split(r'(\*\*.+?\*\*|`.+?`)', text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = par.add_run(tok[1:-1])
            r.font.name = 'Consolas'
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            par.add_run(tok)


def strip_links(text):
    return re.sub(r'(?<!\!)\[([^\]]+)\]\([^)]+\)', r'\1', text)


def build():
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.8)

    lines = open(SRC, encoding='utf-8').read().split('\n')
    i, in_html_comment = 0, False
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        # fenced code block ``` ... ``` → literal monospace, not parsed as markdown
        if line.strip().startswith('```'):
            buf = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i]); i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            r = p.add_run('\n'.join(buf))
            r.font.name = 'Consolas'; r.font.size = Pt(9)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F2F3F5')
            pPr.append(shd)
            continue

        # HTML comment blocks (<!-- ... -->) → skipped
        if '<!--' in line:
            in_html_comment = '-->' not in line
            i += 1
            continue
        if in_html_comment:
            if '-->' in line:
                in_html_comment = False
            i += 1
            continue

        if re.fullmatch(r'-{3,}', line.strip()):
            i += 1
            continue

        # image ![alt](path)
        m = re.match(r'^\s*!\[([^\]]*)\]\(([^)]+)\)', line)
        if m:
            alt, path = m.group(1), m.group(2).strip()
            fp = path if os.path.isabs(path) else os.path.join(ROOT, path)
            if os.path.exists(fp):
                try:
                    doc.add_picture(shrink(fp), width=IMG_W)
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph('[imagem: %s — %s]' % (path, e))
            else:
                p = doc.add_paragraph()
                r = p.add_run('[imagem ausente: %s]' % path)
                r.italic = True
                r.font.color.rgb = GREY
            i += 1
            continue

        # headings
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            lvl = len(m.group(1))
            txt = strip_links(m.group(2))
            hp = doc.add_heading(level=min(lvl, 4))
            r = hp.add_run(txt)
            if lvl == 1:
                r.font.size = Pt(19); r.font.color.rgb = BLUE
            elif lvl == 2:
                r.font.size = Pt(15); r.font.color.rgb = BLUE
            elif lvl == 3:
                r.font.size = Pt(12.5)
            else:
                r.font.size = Pt(11); r.font.color.rgb = GREY
            i += 1
            continue

        # table
        if line.strip().startswith('|'):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl.append(lines[i].strip()); i += 1
            rows = [r for r in tbl if not re.fullmatch(r'\|[\s:|-]+\|', r)]
            cells = [[c.strip() for c in r.strip('|').split('|')] for r in rows]
            if cells:
                ncol = max(len(r) for r in cells)
                t = doc.add_table(rows=0, cols=ncol)
                t.style = 'Light Grid Accent 1'
                for ri, row in enumerate(cells):
                    tr = t.add_row().cells
                    for ci in range(ncol):
                        val = strip_links(row[ci]) if ci < len(row) else ''
                        add_inline(tr[ci].paragraphs[0], val)
                        if ri == 0:
                            shade(tr[ci], '0066CC')
                            for rn in tr[ci].paragraphs[0].runs:
                                rn.bold = True; rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            doc.add_paragraph()
            continue

        # blockquote
        if line.strip().startswith('>'):
            q = line.strip().lstrip('>').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'EAF2FB')
            pPr.append(shd)
            add_inline(p, strip_links(q))
            for rn in p.runs:
                rn.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5a)
            i += 1
            continue

        # checkbox list  - [ ] / - [x]
        m = re.match(r'^\s*[-*]\s+\[([ xX])\]\s+(.*)', line)
        if m:
            done = m.group(1).lower() == 'x'
            p = doc.add_paragraph(style='List Bullet')
            box = p.add_run(('☑ ' if done else '☐ '))
            box.font.name = 'Segoe UI Symbol'
            add_inline(p, strip_links(m.group(2)))
            i += 1
            continue

        # numbered list
        m = re.match(r'^\s*\d+\.\s+(.*)', line)
        if m:
            add_inline(doc.add_paragraph(style='List Number'), strip_links(m.group(1)))
            i += 1
            continue

        # bullet list
        m = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m:
            style = 'List Bullet 2' if len(m.group(1)) >= 2 else 'List Bullet'
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style='List Bullet')
            add_inline(p, strip_links(m.group(2)))
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        add_inline(doc.add_paragraph(), strip_links(line))
        i += 1

    doc.save(OUT)
    imgs = len(doc.inline_shapes)
    print('OK  ->  %s  (%d parágrafos, %d imagens)' % (OUT, len(doc.paragraphs), imgs))


if __name__ == '__main__':
    build()
